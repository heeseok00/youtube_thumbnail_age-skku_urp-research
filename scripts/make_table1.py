from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 페이지 여백 설정 (좁게) ──────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── 스타일 헬퍼 ─────────────────────────────────────────────────────
def set_font(run, size=10, bold=False, italic=False):
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic

def para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT):
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """특정 셀의 테두리 설정"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),  val.get("val",  "single"))
            el.set(qn("w:sz"),   val.get("sz",   "4"))
            el.set(qn("w:space"),"0")
            el.set(qn("w:color"),val.get("color","000000"))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def remove_cell_border(cell, sides=("top","bottom","left","right")):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),  "none")
        el.set(qn("w:sz"),   "0")
        el.set(qn("w:space"),"0")
        el.set(qn("w:color"),"auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)

# ── 표 제목 (표 위에 위치) ───────────────────────────────────────────
title_p = doc.add_paragraph()
para_fmt(title_p, WD_ALIGN_PARAGRAPH.LEFT)
title_p.paragraph_format.space_after = Pt(4)
r1 = title_p.add_run("Table 1. ")
set_font(r1, size=10, bold=True)
r2 = title_p.add_run("Key Discriminating Signals Across Modalities")
set_font(r2, size=10, italic=True)

# ── 표 데이터 ────────────────────────────────────────────────────────
# (modality, feature, ~34, 65+, evidence)
# modality가 빈 문자열이면 이전 행과 병합 대상 (시각적 들여쓰기로 처리)
rows_data = [
    # header
    ("Modality", "Feature / Signal", "~34 group", "65+ group", "Evidence"),
    # visual
    ("Visual",  "text_ratio (MDI=0.393)", "0.234",             "0.345",              "MW p = 2.02e-118"),
    ("",        "person_count",           "1 person (26.4%)",  "3+ persons (39.0%)", "chi² p = 8.26e-7"),
    ("",        "color_saturation",       "77.8",              "92.4",               "MW p = 8.90e-59"),
    ("",        "color_brightness",       "124.6",             "118.4",              "MW p = 6.38e-8"),
    ("",        "color_entropy",          "2.327",             "2.515",              "MW p = 2.39e-25"),
    ("",        "expression_dominant",    "positive / neutral","negative",           "chi² = 66.99"),
    # title
    ("Title",   "key terms (KLUE-RoBERTa + LIME)",  "incident, accident, real, actual","shock, twist, touching, tears","KLUE-RoBERTa + LIME"),
    ("",        "structural signals",      "emoji use↑, CLIP similarity↑","word count↑, noun ratio↑, curiosity↑","XGBoost SHAP"),
    # semantics
    ("Semantics","SHAP tokens",           "painting, sitting, mask","suit, crying, yelling","LLaVA SHAP"),
]

COL_WIDTHS = [Cm(2.3), Cm(4.5), Cm(3.2), Cm(3.2), Cm(3.3)]
THIN  = {"val": "single", "sz": "4",  "color": "000000"}
THICK = {"val": "single", "sz": "12", "color": "000000"}
NONE  = "none"

table = doc.add_table(rows=len(rows_data), cols=5)
table.style = "Light List"

for r_idx, row_vals in enumerate(rows_data):
    row = table.rows[r_idx]
    is_header = (r_idx == 0)
    is_last   = (r_idx == len(rows_data) - 1)

    for c_idx, val in enumerate(row_vals):
        cell = row.cells[c_idx]
        cell.width = COL_WIDTHS[c_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 텍스트
        p = cell.paragraphs[0]
        para_fmt(p)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(val)
        set_font(run, size=10, bold=is_header)

        # 테두리: 기본은 모두 제거
        remove_cell_border(cell)

        # 헤더 행: 위 굵은선 + 아래 얇은선
        if is_header:
            set_cell_border(cell, top=THICK, bottom=THIN)
        # 마지막 행: 아래 굵은선
        elif is_last:
            set_cell_border(cell, bottom=THICK)

# modality 열: 비어있는 셀은 시각적으로 들여쓰기
for r_idx in range(1, len(rows_data)):
    cell = table.rows[r_idx].cells[0]
    p = cell.paragraphs[0]
    if rows_data[r_idx][0] == "":
        p.paragraph_format.left_indent = Cm(0)  # 빈 채로 두면 됨

# ── 표 주석 ──────────────────────────────────────────────────────────
doc.add_paragraph()
note_p = doc.add_paragraph()
para_fmt(note_p)
rn = note_p.add_run("Note. ")
set_font(rn, size=9, bold=True)
rn2 = note_p.add_run(
    "MW = Mann-Whitney U test; chi² = chi-square test; MDI = Mean Decrease in Impurity "
    "(Random Forest feature importance). All reported differences are statistically significant "
    "(p < .001 unless otherwise noted). ~34 group: n = 3,701; 65+ group: n = 3,057. "
    "Visual features extracted via: EasyOCR (text_ratio), YOLOv8 (person_count), "
    "HSV color space (color features), FER (expression_dominant)."
)
set_font(rn2, size=9)

# ── 저장 ─────────────────────────────────────────────────────────────
out_path = "/home/urp_jwl2/26-1_URP/Table1_discriminating_signals.docx"
doc.save(out_path)
print(f"저장 완료: {out_path}")
